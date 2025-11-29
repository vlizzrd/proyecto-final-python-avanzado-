

from docx import Document

def crear_plantilla(path="templates/boleta_template.docx"):
    doc = Document()
    doc.add_heading("Boleta de Calificaciones", level=1)
    doc.add_paragraph("Nombre: {{ nombre }}")
    doc.add_paragraph("ID: {{ id }}")
    doc.add_paragraph("Curso: {{ curso }}")
    doc.add_paragraph("Promedio: {{ promedio }}")
    doc.add_paragraph("Estado: {{ estado }}")
    doc.add_paragraph()
    doc.add_paragraph("Grafica de distribucion del curso:")
    doc.add_paragraph("{{ grafica }}")
    doc.add_paragraph()
    doc.add_paragraph("Comentario: {{ comentario }}")

    doc.save(path)
    print(f"Plantilla creada en: {path}")

if _name_ == "_main_":
    import os
    os.makedirs("templates", exist_ok=True)
    crear_plantilla()


import pandas as pd
import os

def generar_datos(path="datos_calificaciones.xlsx"):

    data = [
        {"id": "A001", "nombre": "Ana Carolina", "email": "ana.carolina@example.com", "tipo": "grado", "nota": 85},
        {"id": "A002", "nombre": "Luis Sarcos", "email": "kuis.sarcos@example", "tipo": "grado", "nota": 58},
        {"id": "A003", "nombre": "Carla Rodriguez", "email": "carla.rodriguez@example.com", "tipo": "postgrado", "nota": 92},
        {"id": "A004", "nombre": "Pedro Soto", "email": "pedro.soto@example.com", "tipo": "postgrado", "nota": 47},
        {"id": "A005", "nombre": "Daniel Espitia", "email": "daniel.espitia@example.com", "tipo": "grado", "nota": 74},
    ]
    df = pd.DataFrame(data)
    df.to_excel(path, index=False)
    print(f"Archivo de ejemplo creado en: {path}")

if _name_ == "_main_":
    generar_datos()